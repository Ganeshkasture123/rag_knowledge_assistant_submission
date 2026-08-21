from pathlib import Path
import hashlib
import re

import pymupdf
from docx import Document
from openpyxl import load_workbook

from .clients import openai_client, search_client
from .config import (
    EMBEDDING_DEPLOYMENT,
    CHUNK_SIZE,
    CHUNK_OVERLAP,
)


SUPPORTED = {".pdf", ".docx", ".xlsx"}


def clean(text) -> str:
    if text is None:
        return ""

    text = str(text).replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


def split_text(
    text: str,
    size=CHUNK_SIZE,
    overlap=CHUNK_OVERLAP,
):
    text = clean(text)

    if not text:
        return []

    chunks = []
    start = 0

    while start < len(text):
        end = min(start + size, len(text))

        if end < len(text):
            boundaries = [
                text.rfind("\n\n", start, end),
                text.rfind(". ", start, end),
                text.rfind("\n", start, end),
            ]

            boundary = max(boundaries)

            if boundary > start + int(size * 0.6):
                end = boundary + (
                    1 if text[boundary:boundary + 2] != "\n\n" else 0
                )

        chunk = text[start:end].strip()

        if chunk:
            chunks.append(chunk)

        if end >= len(text):
            break

        start = max(end - overlap, start + 1)

    return chunks


def parse_pdf(path: Path):
    records = []

    with pymupdf.open(path) as doc:
        for page_no, page in enumerate(doc, 1):
            text = page.get_text("text")

            for section, chunk in enumerate(
                split_text(text), 1
            ):
                records.append(
                    (
                        chunk,
                        {
                            "page": page_no,
                            "section": section,
                        },
                    )
                )

    return records


def parse_docx(path: Path):
    doc = Document(path)
    parts = []

    for paragraph in doc.paragraphs:
        text = clean(paragraph.text)

        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            values = [
                clean(cell.text)
                for cell in row.cells
            ]

            values = [v for v in values if v]

            if values:
                parts.append(" | ".join(values))

    text = "\n".join(parts)

    return [
        (chunk, {"section": i})
        for i, chunk in enumerate(
            split_text(text), 1
        )
    ]


def parse_xlsx(path: Path):
    records = []

    workbook = load_workbook(
        path,
        data_only=True,
    )

    for sheet in workbook.worksheets:
        rows = []

        for row in sheet.iter_rows(
            values_only=True
        ):
            values = [
                clean(value)
                for value in row
                if value is not None
            ]

            values = [v for v in values if v]

            if values:
                rows.append(" | ".join(values))

        text = "\n".join(rows)

        for section, chunk in enumerate(
            split_text(text), 1
        ):
            records.append(
                (
                    chunk,
                    {
                        "sheet": sheet.title,
                        "section": section,
                    },
                )
            )

    return records


def parse_file(path: Path):
    parsers = {
        ".pdf": parse_pdf,
        ".docx": parse_docx,
        ".xlsx": parse_xlsx,
    }

    parser = parsers.get(path.suffix.lower())

    return parser(path) if parser else []


def embed(text: str):
    response = openai_client.embeddings.create(
        model=EMBEDDING_DEPLOYMENT,
        input=text,
    )

    return response.data[0].embedding


def create_document_id(
    source: str,
    chunk_id: int,
):
    return hashlib.sha256(
        f"{source}:{chunk_id}".encode()
    ).hexdigest()


def upload_batches(
    documents,
    batch_size=100,
):
    total = len(documents)

    for start in range(
        0,
        total,
        batch_size,
    ):
        batch = documents[
            start:start + batch_size
        ]

        results = search_client.upload_documents(
            documents=batch
        )

        failed = [
            result
            for result in results
            if not result.succeeded
        ]

        if failed:
            print(
                f"Failed documents: {len(failed)}"
            )

            for item in failed[:5]:
                print(
                    item.key,
                    item.error_message,
                )

        print(
            f"Uploaded "
            f"{min(start + batch_size, total)}/{total}"
        )


def ingest_directory(
    directory="knowledge_base",
):
    root = Path(directory)

    if not root.exists():
        raise FileNotFoundError(
            f"Knowledge base not found: "
            f"{root.resolve()}"
        )

    files = [
        path
        for path in root.rglob("*")
        if (
            path.is_file()
            and path.suffix.lower() in SUPPORTED
        )
    ]

    documents = []

    print(f"Found {len(files)} documents.")

    for path in files:
        print(f"Processing: {path}")

        try:
            source = str(
                path.relative_to(root)
            )

            records = parse_file(path)

            for chunk_id, (
                chunk,
                location,
            ) in enumerate(records):

                documents.append(
                    {
                        "id": create_document_id(
                            source,
                            chunk_id,
                        ),
                        "content": chunk,
                        "title": path.name,
                        "category": path.parent.name,
                        "source": source,
                        "page": int(
                            location.get(
                                "page",
                                0,
                            )
                        ),
                        "sheet": location.get(
                            "sheet",
                            "",
                        ),
                        "chunk_id": chunk_id,
                        "contentVector": embed(
                            chunk
                        ),
                    }
                )

        except Exception as exc:
            print(
                f"ERROR processing {path}: {exc}"
            )

    print(
        f"Created {len(documents)} chunks."
    )

    if documents:
        upload_batches(documents)

    return len(documents)